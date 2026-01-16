"""
CV Text Extractor Module
Extracts text from PDF, DOC, and DOCX files with OCR fallback for scanned documents
"""
import os
import re
from typing import Dict, Optional, List
from io import BytesIO
import tempfile

# Document processing
from docx import Document
import subprocess

# PDF processing
from pdfminer.high_level import extract_text as pdf_extract_text
from pdfminer.pdfpage import PDFPage
from pdf2image import convert_from_path
import pytesseract
from PIL import Image, ImageEnhance, ImageFilter

# Configure Tesseract path
try:
    from config import Config
    if hasattr(Config, 'TESSERACT_CMD') and os.path.exists(Config.TESSERACT_CMD):
        pytesseract.pytesseract.tesseract_cmd = Config.TESSERACT_CMD
    POPPLER_PATH = getattr(Config, 'POPPLER_PATH', None)
except:
    # Fallback: try common locations
    if os.path.exists(r'C:\Program Files\Tesseract-OCR\tesseract.exe'):
        pytesseract.pytesseract.tesseract_cmd = r'C:\Program Files\Tesseract-OCR\tesseract.exe'
    POPPLER_PATH = r'C:\poppler\poppler-24.02.0\Library\bin' if os.path.exists(r'C:\poppler\poppler-24.02.0\Library\bin') else None


class CVTextExtractor:
    """Extract and parse text from CV documents"""
    
    def __init__(self):
        self.min_text_length = 50  # Minimum characters for valid extraction
        
    def extract_from_file(self, filepath: str) -> Dict:
        """
        Main entry point for CV text extraction
        
        Args:
            filepath: Path to the CV file
            
        Returns:
            Dict with 'full_text' and 'parsed_fields'
        """
        ext = os.path.splitext(filepath)[1].lower()
        
        if ext == '.pdf':
            return self._extract_from_pdf(filepath)
        elif ext == '.docx':
            return self._extract_from_docx(filepath)
        elif ext == '.doc':
            return self._extract_from_doc(filepath)
        else:
            raise ValueError(f"Unsupported file format: {ext}")
    
    def _extract_from_docx(self, filepath: str) -> Dict:
        """Extract text from DOCX file"""
        try:
            doc = Document(filepath)
            
            # Extract all paragraphs
            paragraphs = [para.text for para in doc.paragraphs if para.text.strip()]
            
            # Also extract text from tables (if any)
            for table in doc.tables:
                for row in table.rows:
                    for cell in row.cells:
                        if cell.text.strip():
                            paragraphs.append(cell.text.strip())
            
            full_text = '\n'.join(paragraphs)
            
            # Count pages approximately (assuming ~500 chars per page)
            approx_pages = max(1, len(full_text) // 500)
            print(f"DOCX extraction: {len(paragraphs)} paragraphs, {len(full_text)} characters (approx {approx_pages} pages)")
            
            # Debug: Print first 500 chars to see structure
            print(f"First 500 chars of extracted text:\\n{full_text[:500]}...")
            
            return {
                'full_text': full_text,
                'parsed_fields': self._parse_cv_text(full_text),
                'extraction_method': 'docx'
            }
        except Exception as e:
            raise Exception(f"Error extracting DOCX: {str(e)}")
    
    def _extract_from_doc(self, filepath: str) -> Dict:
        """Convert DOC to DOCX then extract"""
        try:
            # Convert DOC to DOCX using LibreOffice (requires LibreOffice installed)
            # Alternative: use subprocess to call 'antiword' or similar tool
            output_dir = tempfile.mkdtemp()
            output_file = os.path.join(output_dir, 'converted.docx')
            
            # Try using LibreOffice for conversion
            try:
                subprocess.run([
                    'soffice',
                    '--headless',
                    '--convert-to', 'docx',
                    '--outdir', output_dir,
                    filepath
                ], check=True, timeout=30)
                
                # Find the converted file
                converted_files = [f for f in os.listdir(output_dir) if f.endswith('.docx')]
                if converted_files:
                    converted_path = os.path.join(output_dir, converted_files[0])
                    return self._extract_from_docx(converted_path)
            except (subprocess.CalledProcessError, FileNotFoundError):
                # LibreOffice not available, try alternative method
                pass
            
            # Fallback: Try using antiword if available
            try:
                result = subprocess.run(
                    ['antiword', filepath],
                    capture_output=True,
                    text=True,
                    timeout=30
                )
                if result.returncode == 0:
                    full_text = result.stdout
                    return {
                        'full_text': full_text,
                        'parsed_fields': self._parse_cv_text(full_text),
                        'extraction_method': 'antiword'
                    }
            except FileNotFoundError:
                pass
            
            raise Exception("No DOC converter available. Please install LibreOffice or antiword.")
            
        except Exception as e:
            raise Exception(f"Error extracting DOC: {str(e)}")
    
    def _extract_from_pdf(self, filepath: str) -> Dict:
        """Extract text from PDF with OCR fallback"""
        try:
            # Try normal text extraction first - extract_text processes ALL pages by default
            full_text = pdf_extract_text(filepath)
            
            # Clean up extracted text
            full_text = full_text.strip()
            
            print(f"PDF text extraction: {len(full_text)} characters extracted")
            
            # Check if extraction was successful
            # Lower threshold - even 20 chars might indicate text-based PDF
            if len(full_text) >= 20:
                # Check if the text has meaningful content (not just spaces/special chars)
                meaningful_chars = sum(1 for c in full_text if c.isalnum())
                
                if meaningful_chars >= 15:  # At least 15 alphanumeric characters
                    print(f"PDF has extractable text ({meaningful_chars} meaningful chars)")
                    return {
                        'full_text': full_text,
                        'parsed_fields': self._parse_cv_text(full_text),
                        'extraction_method': 'pdfminer'
                    }
            
            # Insufficient text - likely a scanned PDF, apply OCR
            print(f"Extracted text insufficient ({len(full_text)} chars), applying OCR...")
            ocr_text = self._ocr_pdf(filepath)
            
            # Use OCR text if it's more substantial than extracted text
            if len(ocr_text.strip()) > len(full_text):
                final_text = ocr_text
            else:
                # Merge both if extracted text has some content
                final_text = (full_text + '\n' + ocr_text).strip()
            
            return {
                'full_text': final_text,
                'parsed_fields': self._parse_cv_text(final_text),
                'extraction_method': 'ocr' if not full_text else 'pdfminer+ocr'
            }
            
        except Exception as e:
            # If pdfminer fails, try OCR directly
            print(f"PDFMiner error: {str(e)}, falling back to OCR")
            try:
                ocr_text = self._ocr_pdf(filepath)
                return {
                    'full_text': ocr_text,
                    'parsed_fields': self._parse_cv_text(ocr_text),
                    'extraction_method': 'ocr_only'
                }
            except Exception as ocr_error:
                raise Exception(f"Error extracting PDF: {str(e)}. OCR also failed: {str(ocr_error)}")
    
    def _ocr_pdf(self, filepath: str) -> str:
        """Apply OCR to PDF pages with image preprocessing"""
        try:
            # Convert PDF pages to images at high DPI for better OCR accuracy
            print(f"Converting PDF to images for OCR...")
            
            # Use poppler path if available
            if POPPLER_PATH:
                images = convert_from_path(filepath, dpi=300, fmt='png', poppler_path=POPPLER_PATH)
            else:
                images = convert_from_path(filepath, dpi=300, fmt='png')
            
            print(f"OCR processing {len(images)} pages from PDF...")
            
            ocr_texts = []
            for i, image in enumerate(images):
                # Preprocess image for better OCR
                # Convert to grayscale
                image = image.convert('L')
                
                # Enhance contrast
                enhancer = ImageEnhance.Contrast(image)
                image = enhancer.enhance(2.0)
                
                # Sharpen image
                image = image.filter(ImageFilter.SHARPEN)
                
                # Apply OCR to each page with custom config
                # PSM 3 = Fully automatic page segmentation, but no OSD (default)
                # PSM 6 = Assume a single uniform block of text
                custom_config = r'--oem 3 --psm 3'
                text = pytesseract.image_to_string(image, lang='eng', config=custom_config)
                
                print(f"Page {i+1}: {len(text)} characters extracted via OCR")
                
                # Only include pages with meaningful text
                if len(text.strip()) > 10:  # Lower threshold - even 10 chars could be useful
                    ocr_texts.append(text)
                else:
                    print(f"Page {i+1}: Skipped (insufficient text)")
            
            if not ocr_texts:
                print("WARNING: No text extracted from any page via OCR")
                return ''
            
            combined = '\n\n'.join(ocr_texts)
            print(f"Total OCR text: {len(combined)} characters from {len(ocr_texts)} pages")
            return combined
            
        except Exception as e:
            print(f"OCR Error details: {str(e)}")
            raise Exception(f"Error during OCR: {str(e)}. Make sure Tesseract is installed and in PATH.")
    
    def _parse_cv_text(self, text: str) -> Dict:
        """
        Parse CV text into structured fields
        
        Returns:
            Dict with extracted fields: name, email, phone, skills, education, experience, languages, certifications
        """
        parsed = {
            'name': self._extract_name(text),
            'email': self._extract_email(text),
            'phone': self._extract_phone(text),
            'location': self._extract_location(text),
            'skills': self._extract_skills(text),
            'education': self._extract_education(text),
            'experience': self._extract_experience(text),
            'summary': self._extract_summary(text),
            'languages': self._extract_languages(text),
            'certifications': self._extract_certifications(text)
        }
        
        return parsed
    
    def _extract_name(self, text: str) -> Optional[str]:
        """Extract candidate name - handles both formatted and jumbled text"""
        
        # Method 1: Try finding name after common keywords
        name_patterns = [
            r'(?:name|full name|candidate):\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3})',
        ]
        
        for pattern in name_patterns:
            match = re.search(pattern, text, re.IGNORECASE)
            if match:
                name = match.group(1).strip()
                if 2 <= len(name.split()) <= 4:
                    print(f"Name found via pattern: {name}")
                    return name
        
        # Method 2: Try first non-empty lines (for formatted CVs)
        lines = text.strip().split('\n')
        for line in lines[:5]:  # Check first 5 lines
            line = line.strip()
            if line and len(line) < 50 and not any(kw in line.lower() for kw in ['email', 'phone', 'address', 'linkedin', 'profile', 'skill']):
                # Basic name validation - 2-4 words, capitalized
                words = line.split()
                if 2 <= len(words) <= 4 and all(w and w[0].isupper() for w in words if w and w[0].isalpha()):
                    print(f"Name found in first lines: {line}")
                    return line
        
        # Method 3: For jumbled text, look for pattern of 2-4 capitalized words near the start
        # Extract first 200 chars
        text_start = text[:200]
        
        # Look for sequences of 2-4 capitalized words
        name_pattern = r'\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+){1,3})\b'
        matches = re.findall(name_pattern, text_start)
        
        for match in matches:
            words = match.split()
            # Verify it's not a common section header or location
            if len(words) >= 2 and not any(kw in match.lower() for kw in [
                'profile', 'summary', 'objective', 'skills', 'education', 'experience',
                'technical', 'professional', 'work', 'employment', 'university', 'college',
                'data science', 'software engineer', 'royal university'
            ]):
                # Prefer matches that come before email/phone
                email_pos = text_start.lower().find('@')
                match_pos = text_start.find(match)
                if email_pos == -1 or match_pos < email_pos:
                    print(f"Name found via capitalized words: {match}")
                    return match
        
        return None
    
    def _extract_email(self, text: str) -> Optional[str]:
        """Extract email address"""
        email_pattern = r'\b[A-Za-z0-9._%+-]+@[A-Za-z0-9.-]+\.[A-Z|a-z]{2,}\b'
        matches = re.findall(email_pattern, text)
        return matches[0] if matches else None
    
    def _extract_phone(self, text: str) -> Optional[str]:
        """Extract phone number - handles various formats and jumbled text"""
        # Match various phone formats
        phone_patterns = [
            r'(?:Phone|Tel|Mobile|Contact):\s*([\d\s\-\.\(\)\+]+)',  # With label
            r'\+?\d{1,3}[\s.\-]?\(?\d{3}\)?[\s.\-]?\d{3}[\s.\-]?\d{3,4}',  # International format
            r'\b0\d{2}[\s.\-]?\d{3}[\s.\-]?\d{3,4}\b',  # Format like 099 722 116
            r'\b0\d{8,10}\b',  # Format like 0997221160
            r'\b\d{3}[\s.\-]\d{3}[\s.\-]\d{3,4}\b',  # Format like 099 722 116
            r'\+\d{1,3}[\s.\-]?\d{1,4}[\s.\-]?\d{1,4}[\s.\-]?\d{1,4}[\s.\-]?\d{1,4}',  # International with +
        ]
        
        for pattern in phone_patterns:
            matches = re.findall(pattern, text)
            if matches:
                phone = matches[0] if isinstance(matches[0], str) else ''.join(matches[0])
                phone = phone.strip()
                
                # Clean up: remove excessive spaces
                phone = re.sub(r'\s+', ' ', phone)
                
                # Validate: should have at least 8 digits
                digits_only = re.sub(r'\D', '', phone)
                if len(digits_only) >= 8:
                    print(f"Phone found: {phone} (digits: {digits_only})")
                    return phone
        
        return None
    
    def _extract_location(self, text: str) -> Optional[str]:
        """Extract location/address"""
        # Look for common location patterns
        location_patterns = [
            r'Address:\s*([^\n]+)',
            r'Location:\s*([^\n]+)',
            r'\b([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?),\s*([A-Z][a-z]+(?:\s+[A-Z][a-z]+)?)(?:,\s*([A-Z][a-z]+))?\b',  # City, State/Province, Country
        ]
        
        for pattern in location_patterns:
            matches = re.search(pattern, text, re.IGNORECASE)
            if matches:
                if matches.lastindex == 1:
                    return matches.group(1).strip()
                else:
                    # Multi-group match (city, state, country)
                    location_parts = [g for g in matches.groups() if g]
                    return ', '.join(location_parts)
        
        return None
    
    def _extract_skills(self, text: str) -> List[str]:
        """Extract skills section - handles both formatted and jumbled text"""
        skills = []
        
        # Look for skills section
        skills_section = self._extract_section(text, ['skills', 'technical skills', 'core competencies', 'key skills', 'competencies', 'hard skills', 'soft skills'])
        
        if skills_section:
            print(f"Skills section found: {len(skills_section)} characters")
            
            # Common technical skills to look for
            common_skills = [
                # Programming Languages
                'python', 'java', 'javascript', 'typescript', 'c++', 'c#', 'ruby', 'php', 'swift', 
                'kotlin', 'go', 'rust', 'r', 'matlab', 'scala', 'perl', 'sql',
                # Web Technologies
                'html', 'css', 'react', 'angular', 'vue', 'node.js', 'express', 'django', 'flask',
                'spring', 'asp.net', 'laravel', 'jquery', 'bootstrap', 'tailwind',
                # Databases
                'mysql', 'postgresql', 'mongodb', 'redis', 'oracle', 'sqlite', 'cassandra',
                'dynamodb', 'firebase',
                # Cloud & DevOps
                'aws', 'azure', 'gcp', 'docker', 'kubernetes', 'jenkins', 'git', 'github', 'gitlab',
                'terraform', 'ansible', 'ci/cd',
                # Data Science & ML
                'machine learning', 'deep learning', 'tensorflow', 'pytorch', 'scikit-learn',
                'pandas', 'numpy', 'jupyter', 'data analysis', 'data visualization',
                # Tools & Software
                'excel', 'powerpoint', 'word', 'outlook', 'photoshop', 'illustrator',
                'figma', 'sketch', 'jira', 'confluence', 'slack', 'trello',
                # Soft Skills
                'communication', 'leadership', 'teamwork', 'problem-solving', 'time management',
                'critical thinking', 'analytical', 'creativity', 'adaptability', 'collaboration'
            ]
            
            # Filter out section headers that might be included
            section_headers = ['language', 'languages', 'soft skills', 'hard skills', 'technical skills', 
                             'certifications', 'education', 'experience', 'projects', 'achievements']
            
            # Method 1: Split by common delimiters (comma, bullet points, pipe, newline, semicolon)
            items = re.split(r'[,•·\|\/\n;]', skills_section)
            for item in items:
                item = item.strip()
                # Clean up bullet characters and numbers
                item = re.sub(r'^[\d\.\)\-\*\•\·\s]+', '', item)
                item = item.strip()
                
                # Skip if it's a section header
                if item.lower() in section_headers:
                    continue
                
                # Skip if too short or too long
                if len(item) < 2 or len(item) > 50:
                    continue
                
                # Skip if it contains too many numbers (likely not a skill)
                if sum(c.isdigit() for c in item) > len(item) * 0.5:
                    continue
                
                # Add if it looks like a skill
                if item and item not in skills:
                    skills.append(item)
            
            # Method 2: Word-by-word matching for jumbled text
            if len(skills) < 5:  # If we didn't find many skills, try harder
                words_lower = skills_section.lower()
                for skill in common_skills:
                    # Look for the skill as a whole word
                    if re.search(r'\b' + re.escape(skill) + r'\b', words_lower):
                        # Capitalize properly
                        skill_formatted = skill.upper() if skill.isupper() or len(skill) <= 3 else skill.title()
                        if skill_formatted not in skills and skill not in skills:
                            skills.append(skill_formatted)
                            print(f"Found common skill: {skill_formatted}")
        
        print(f"Extracted {len(skills)} skills")
        return skills[:100]  # Limit to 100 skills
    
    def _extract_education(self, text: str) -> List[Dict]:
        """Extract education history"""
        education = []
        seen_entries = set()
        
        edu_section = self._extract_section(text, ['education', 'academic background', 'qualifications'])
        
        if edu_section:
            print(f"Education section found: {len(edu_section)} characters")
            
            # Split by double newlines first to get distinct entries
            entries = re.split(r'\n\s*\n', edu_section)
            
            for entry in entries:
                entry = entry.strip()
                if not entry or len(entry) < 10:
                    continue
                
                # Look for degree keywords
                if any(keyword in entry.lower() for keyword in ['bachelor', 'master', 'phd', 'degree', 'diploma', 'university', 'college', 'high school']):
                    # Clean up the entry
                    entry_clean = re.sub(r'^[\d\.\)\-\*\•\·\s]+', '', entry)
                    entry_clean = ' '.join(entry_clean.split())  # Normalize whitespace
                    
                    # Avoid duplicates
                    if entry_clean and entry_clean not in seen_entries:
                        seen_entries.add(entry_clean)
                        
                        # Parse structured fields
                        parsed_edu = self._parse_education_entry(entry_clean)
                        education.append(parsed_edu)
        
        print(f"Extracted {len(education)} education entries")
        return education  # No artificial limit - return all entries
    
    def _parse_education_entry(self, entry):
        """Parse education entry into structured fields: degree, institution, year, description"""
        degree = ''
        institution = ''
        year = ''
        description = entry
        
        lines = entry.split('\n') if '\n' in entry else [entry]
        first_line = lines[0] if lines else entry
        
        # Extract year
        year_patterns = [
            r'\b(\d{4}\s*[-–—]\s*(?:\d{4}|present|current))\b',  # 2020-2024, 2023-Present
            r'\b(\d{4})\b(?!\s*[-–—])',  # Single year like 2024
        ]
        for pattern in year_patterns:
            match = re.search(pattern, first_line, re.IGNORECASE)
            if match:
                year = match.group(1).strip()
                first_line = first_line.replace(match.group(0), '').strip()
                break
        
        # Extract degree and institution
        # Pattern 1: "Degree at Institution"
        if ' at ' in first_line.lower():
            parts = re.split(r'\s+at\s+', first_line, maxsplit=1, flags=re.IGNORECASE)
            if len(parts) == 2:
                degree = parts[0].strip()
                institution = parts[1].strip()
        
        # Pattern 2: "Degree, Institution" or "Degree,Institution"
        elif ',' in first_line:
            parts = first_line.split(',', 1)
            if len(parts) == 2:
                degree = parts[0].strip()
                institution = parts[1].strip()
        
        # Pattern 3: "Degree - Institution" or "Degree – Institution"
        elif re.search(r'\s+[-–—]\s+', first_line):
            parts = re.split(r'\s+[-–—]\s+', first_line, maxsplit=1)
            if len(parts) == 2:
                degree = parts[0].strip()
                institution = parts[1].strip()
        
        # Pattern 4: Look for degree keywords and institution keywords separately
        else:
            # Find degree (Bachelor, Master, PhD, etc.)
            degree_pattern = r'(bachelor|master|phd|doctorate|diploma|degree|associate|certificate)[\s\w\-]+'
            degree_match = re.search(degree_pattern, first_line, re.IGNORECASE)
            if degree_match:
                degree = degree_match.group(0).strip()
                # Rest is institution
                institution = first_line.replace(degree, '').strip()
            else:
                # Use first line as degree
                degree = first_line.strip()
        
        # Description is the rest of the lines
        if len(lines) > 1:
            description = ' '.join(lines[1:]).strip()
        else:
            description = ''
        
        return {
            'degree': degree[:200] if degree else '',
            'institution': institution[:200] if institution else '',
            'year': year[:100] if year else '',
            'description': description[:300] if description else ''
        }
    
    def _extract_experience(self, text: str) -> List[Dict]:
        """Extract work experience"""
        experience = []
        seen_entries = set()
        
        exp_section = self._extract_section(text, ['experience', 'work experience', 'employment history', 'professional experience', 'volunteer experience'])
        
        if exp_section:
            print(f"Experience section found: {len(exp_section)} characters")
            
            # Split by double newlines or lines that look like job titles (Title, Company pattern)
            entries = re.split(r'\n\s*\n', exp_section)
            
            for entry in entries:
                entry = entry.strip()
                if not entry or len(entry) < 20:
                    continue
                
                # Clean up entry
                entry_clean = re.sub(r'^[\d\.\)\-\*\•\·\s]+', '', entry)
                entry_clean = ' '.join(entry_clean.split())  # Normalize whitespace
                
                # Check if it looks like an experience entry (has dates or job-related keywords)
                if re.search(r'\b(20\d{2}|19\d{2}|current|present)\b', entry_clean.lower()) or \
                   any(kw in entry_clean.lower() for kw in ['intern', 'manager', 'engineer', 'developer', 'coordinator', 'analyst']):
                    
                    # Avoid duplicates
                    if entry_clean and entry_clean not in seen_entries:
                        seen_entries.add(entry_clean)
                        
                        # Parse structured fields
                        parsed_exp = self._parse_experience_entry(entry_clean)
                        experience.append(parsed_exp)
        
        print(f"Extracted {len(experience)} experience entries")
        return experience  # No artificial limit - return all entries
    
    def _parse_experience_entry(self, entry):
        """Parse experience entry into structured fields: title, company, duration, description"""
        title = ''
        company = ''
        duration = ''
        description = entry
        
        lines = entry.split('\n') if '\n' in entry else [entry]
        first_line = lines[0] if lines else entry
        
        # Extract duration (years)
        duration_patterns = [
            r'\b(\d{4}\s*[-–—]\s*(?:\d{4}|present|current))\b',  # 2020-2023, 2025-Current
            r'\b(\d{4})\b(?!\s*[-–—])',  # Single year like 2025
        ]
        for pattern in duration_patterns:
            match = re.search(pattern, first_line, re.IGNORECASE)
            if match:
                duration = match.group(1).strip()
                first_line = first_line.replace(match.group(0), '').strip()
                break
        
        # Extract title and company from first line
        # Pattern 1: "Title, Company" or "Title,Company"
        if ',' in first_line:
            parts = first_line.split(',', 1)
            if len(parts) == 2:
                title = parts[0].strip()
                company = parts[1].strip()
        
        # Pattern 2: "Title at Company"
        elif ' at ' in first_line.lower():
            parts = re.split(r'\s+at\s+', first_line, maxsplit=1, flags=re.IGNORECASE)
            if len(parts) == 2:
                title = parts[0].strip()
                company = parts[1].strip()
        
        # Pattern 3: "Title - Company" or "Title – Company"
        elif re.search(r'\s+[-–—]\s+', first_line):
            parts = re.split(r'\s+[-–—]\s+', first_line, maxsplit=1)
            if len(parts) == 2:
                title = parts[0].strip()
                company = parts[1].strip()
        
        # If no pattern matched, use first line as title
        else:
            title = first_line.strip()
        
        # Description is the rest of the lines
        if len(lines) > 1:
            description = ' '.join(lines[1:]).strip()
        else:
            description = ''
        
        return {
            'title': title[:200] if title else '',
            'company': company[:200] if company else '',
            'duration': duration[:100] if duration else '',
            'description': description[:500] if description else ''
        }
    
    def _extract_summary(self, text: str) -> Optional[str]:
        """Extract professional summary"""
        summary_section = self._extract_section(text, ['summary', 'profile', 'objective', 'about me', 'professional summary'])
        return summary_section[:500] if summary_section else None
    
    def _extract_languages(self, text: str) -> List[Dict]:
        """Extract languages section - handles various formats"""
        languages = []
        seen_languages = set()
        
        lang_section = self._extract_section(text, ['languages', 'language skills', 'language proficiency', 'language'])
        
        if lang_section:
            print(f"Language section found: {len(lang_section)} characters")
            print(f"Language section content: {lang_section[:200]}...")
            
            # Common languages to look for (expanded list)
            common_langs = [
                'english', 'spanish', 'french', 'german', 'chinese', 'mandarin', 'japanese', 'korean', 
                'arabic', 'portuguese', 'russian', 'italian', 'dutch', 'hindi', 'khmer', 'cambodian',
                'vietnamese', 'thai', 'indonesian', 'malay', 'tagalog', 'turkish', 'polish', 'ukrainian',
                'greek', 'czech', 'swedish', 'finnish', 'danish', 'norwegian', 'hebrew', 'bengali',
                'urdu', 'punjabi', 'tamil', 'telugu', 'burmese', 'lao', 'nepali'
            ]
            
            # Proficiency levels
            proficiency_map = {
                'native': 'Native',
                'mother tongue': 'Native',
                'fluent': 'Native',
                'advanced': 'Advanced',
                'proficient': 'Advanced',
                'upper intermediate': 'Advanced',
                'intermediate': 'Intermediate',
                'conversational': 'Intermediate',
                'basic': 'Basic',
                'beginner': 'Basic',
                'elementary': 'Basic',
                'limited': 'Basic',
                'working proficiency': 'Intermediate'
            }
            
            # Method 1: Try comma-separated format (e.g., "English, Khmer, Chinese")
            if ',' in lang_section:
                # Split by commas and process each item
                items = re.split(r'[,;]', lang_section)
                for item in items:
                    item_lower = item.lower().strip()
                    # Remove bullet points, numbers, dashes
                    item_lower = re.sub(r'^[\d\.\)\-\*•·\s]+', '', item_lower).strip()
                    
                    if not item_lower or len(item_lower) > 50:
                        continue
                    
                    # Check if it contains a language name
                    found_lang = None
                    for lang in common_langs:
                        if re.search(r'\b' + re.escape(lang) + r'\b', item_lower):
                            found_lang = lang
                            break
                    
                    if found_lang and found_lang not in seen_languages:
                        # Determine proficiency from the item
                        proficiency = 'Intermediate'  # Default
                        for key, value in proficiency_map.items():
                            if key in item_lower:
                                proficiency = value
                                break
                        
                        # Handle special language names
                        display_name = found_lang.capitalize()
                        if found_lang == 'mandarin':
                            display_name = 'Chinese (Mandarin)'
                        elif found_lang == 'cambodian':
                            display_name = 'Khmer'
                        
                        languages.append({
                            'language': display_name,
                            'proficiency': proficiency
                        })
                        seen_languages.add(found_lang)
                        print(f"Parsed language (comma-sep): {display_name} - {proficiency}")
            
            # Method 2: Line-by-line parsing
            lines = lang_section.split('\n')
            for line in lines:
                line_lower = line.lower().strip()
                if not line_lower or len(line_lower) > 100:
                    continue
                
                # Clean up line
                line_lower = re.sub(r'^[\d\.\)\-\*•·\s]+', '', line_lower).strip()
                
                # Try to parse "Language: Proficiency" format
                colon_match = re.match(r'^\s*([a-z\s]+)\s*[:|-]\s*(.+)$', line_lower)
                if colon_match:
                    lang_name = colon_match.group(1).strip()
                    proficiency_text = colon_match.group(2).strip()
                    
                    # Check if it's a valid language
                    found = False
                    for lang in common_langs:
                        if lang in lang_name and lang not in seen_languages:
                            # Determine proficiency level from the text
                            proficiency = 'Intermediate'  # Default
                            for key, value in proficiency_map.items():
                                if key in proficiency_text:
                                    proficiency = value
                                    break
                            
                            display_name = lang.capitalize()
                            if lang == 'cambodian':
                                display_name = 'Khmer'
                            
                            languages.append({
                                'language': display_name,
                                'proficiency': proficiency
                            })
                            seen_languages.add(lang)
                            print(f"Parsed language (colon format): {display_name} - {proficiency}")
                            found = True
                            break
                    
                    if found:
                        continue
                
                # Method 3: Just check if line contains a common language
                for lang in common_langs:
                    if re.search(r'\b' + re.escape(lang) + r'\b', line_lower) and lang not in seen_languages:
                        # Determine proficiency level
                        proficiency = 'Intermediate'  # Default
                        for key, value in proficiency_map.items():
                            if key in line_lower:
                                proficiency = value
                                break
                        
                        display_name = lang.capitalize()
                        if lang == 'cambodian':
                            display_name = 'Khmer'
                        
                        languages.append({
                            'language': display_name,
                            'proficiency': proficiency
                        })
                        seen_languages.add(lang)
                        print(f"Parsed language (inline): {display_name} - {proficiency}")
                        break
        
        print(f"Extracted {len(languages)} languages")
        return languages  # No limit
    
    def _extract_certifications(self, text: str) -> List[Dict]:
        """Extract certifications section"""
        certifications = []
        
        cert_section = self._extract_section(text, ['certifications', 'certificates', 'licenses', 'professional certifications'])
        
        if cert_section:
            # Split by lines
            lines = cert_section.split('\n')
            
            for line in lines:
                line = line.strip()
                if not line or len(line) < 5:
                    continue
                
                # Skip lines that are just bullet points or numbers
                if re.match(r'^[•·\-\*\d\.]+\s*$', line):
                    continue
                
                # Extract year if present
                year_match = re.search(r'\b(19|20)\d{2}\b', line)
                year = year_match.group(0) if year_match else ''
                
                # Extract organization (words in title case or uppercase)
                org_match = re.search(r'\b([A-Z][A-Za-z]*(?:\s+[A-Z][A-Za-z]*)+)\b', line)
                organization = org_match.group(0) if org_match else ''
                
                certifications.append({
                    'name': line[:200],  # Full line as cert name
                    'organization': organization[:100],
                    'year': year
                })
        
        print(f"Extracted {len(certifications)} certifications")
        return certifications  # No limit
    
    def _extract_section(self, text: str, keywords: List[str]) -> Optional[str]:
        """Extract a section of text based on keywords - handles both proper formatting and jumbled text"""
        text_lower = text.lower()
        
        for keyword in keywords:
            # Pattern 1: Properly formatted with newlines
            pattern_newline = r'\n\s*' + re.escape(keyword) + r'\s*[:|\n]'
            match = re.search(pattern_newline, text_lower)
            
            # Pattern 2: Jumbled text (keyword in uppercase or title case without newlines)
            if not match:
                # Look for keyword in uppercase (e.g., "SKILLS" or "EXPERIENCE")
                pattern_upper = re.escape(keyword.upper()) + r'(?:\s*[:|\s])'
                match = re.search(pattern_upper, text_lower)
            
            # Pattern 3: Title case in jumbled text (e.g., "Skills" or "Experience")  
            if not match:
                pattern_title = r'\b' + re.escape(keyword.title()) + r'(?:\s*[:|\s])'
                match = re.search(pattern_title, text)
            
            if match:
                start_idx = match.end()
                
                # Find next section - check for any common section headers
                next_section_keywords = [
                    'education', 'experience', 'work experience', 'employment',
                    'skills', 'technical skills', 'core competencies', 'competencies',
                    'certifications', 'certificates', 'certification',
                    'language', 'languages',
                    'projects', 'project',
                    'references', 'volunteer', 'volunteering',
                    'achievements', 'achievement', 'awards', 'award',
                    'soft skills', 'hard skills', 'technical', 'professional',
                    'summary', 'profile', 'objective',
                    'contact', 'personal information'
                ]
                
                # Try different patterns to find end of section
                end_idx = None
                
                # Pattern 1: Newline + section header
                next_pattern_newline = r'\n\s*(?:' + '|'.join([re.escape(kw) for kw in next_section_keywords]) + r')\s*[:|\n]'
                next_match = re.search(next_pattern_newline, text_lower[start_idx:])
                
                if next_match:
                    end_idx = start_idx + next_match.start()
                
                # Pattern 2: Uppercase section header in jumbled text
                if not end_idx:
                    next_pattern_upper = r'(?:' + '|'.join([re.escape(kw.upper()) for kw in next_section_keywords]) + r')(?:\s*[:|\s])'
                    next_match = re.search(next_pattern_upper, text_lower[start_idx:])
                    if next_match:
                        end_idx = start_idx + next_match.start()
                
                # Pattern 3: Title case section header in jumbled text
                if not end_idx:
                    next_pattern_title = r'\b(?:' + '|'.join([re.escape(kw.title()) for kw in next_section_keywords]) + r')(?:\s*[:|\s])'
                    next_match = re.search(next_pattern_title, text[start_idx:])
                    if next_match:
                        end_idx = start_idx + next_match.start()
                
                if end_idx:
                    section_text = text[start_idx:end_idx].strip()
                    next_section_name = text_lower[end_idx:end_idx + 30].strip()
                    print(f"Section '{keyword}' found: {len(section_text)} chars, ends before '{next_section_name[:20]}'")
                else:
                    # Return rest of document (limit to reasonable length to avoid including unrelated content)
                    section_text = text[start_idx:start_idx + 2000].strip()
                    print(f"Section '{keyword}' found: {len(section_text)} chars (limited from rest)")
                
                return section_text
        
        return None


# Singleton instance
extractor = CVTextExtractor()


def extract_cv_text(filepath: str) -> Dict:
    """
    Main function to extract and parse CV text
    
    Args:
        filepath: Path to CV file
        
    Returns:
        Dict with 'full_text' and 'parsed_fields'
    """
    return extractor.extract_from_file(filepath)
